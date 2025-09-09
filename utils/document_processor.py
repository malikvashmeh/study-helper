"""
Document Processing Utilities
Handles PDF, text, and other document formats with enhanced memory management
"""

import os
import logging
import hashlib
import uuid
from datetime import datetime
from typing import List, Dict, Any, Optional
from pathlib import Path

# Document processing imports
import PyPDF2
from docx import Document
import tiktoken

# LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles document loading, text extraction, and memory management"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        # Track processed documents to prevent duplicates
        self.processed_documents = {}
    
    def generate_document_id(self, file_path: str, content: str = None) -> str:
        """Generate unique document ID based on filename and content hash"""
        try:
            file_name = Path(file_path).name
            
            if content:
                # Create hash from content for uniqueness
                content_hash = hashlib.md5(content.encode('utf-8')).hexdigest()[:8]
                return f"{file_name}_{content_hash}"
            else:
                # Fallback to timestamp-based ID
                timestamp = int(datetime.now().timestamp())
                return f"{file_name}_{timestamp}"
                
        except Exception as e:
            logger.warning(f"Error generating document ID: {e}")
            return str(uuid.uuid4())[:8]
    
    def calculate_content_hash(self, content: str) -> str:
        """Calculate MD5 hash of content for duplicate detection"""
        return hashlib.md5(content.encode('utf-8')).hexdigest()
    
    def is_duplicate_document(self, content_hash: str) -> bool:
        """Check if document with this content hash was already processed"""
        return content_hash in self.processed_documents
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {e}")
            raise
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except Exception as e:
            logger.error(f"Error extracting text from TXT {file_path}: {e}")
            raise
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from various file formats"""
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def chunk_text(self, text: str, metadata: Dict[str, Any] = None) -> List[LangChainDocument]:
        """Split text into chunks and return as LangChain Documents with enhanced metadata"""
        if metadata is None:
            metadata = {}
        
        # Add content hash for tracking
        content_hash = self.calculate_content_hash(text)
        metadata['content_hash'] = content_hash
        metadata['processed_timestamp'] = datetime.now().isoformat()
        
        # Split the text
        chunks = self.text_splitter.split_text(text)
        
        # Convert to LangChain Documents
        documents = []
        for i, chunk in enumerate(chunks):
            doc_metadata = metadata.copy()
            doc_metadata.update({
                'chunk_id': i,
                'total_chunks': len(chunks),
                'chunk_hash': self.calculate_content_hash(chunk)
            })
            documents.append(LangChainDocument(
                page_content=chunk,
                metadata=doc_metadata
            ))
        
        return documents
    
    def process_document(self, file_path: str, force_reprocess: bool = False, original_filename: str = None) -> List[LangChainDocument]:
        """Complete document processing pipeline with duplicate detection"""
        try:
            # Extract text
            text = self.extract_text(file_path)
            
            # Check for duplicates
            content_hash = self.calculate_content_hash(text)
            if not force_reprocess and self.is_duplicate_document(content_hash):
                display_name = original_filename or Path(file_path).name
                logger.info(f"Document {display_name} already processed, skipping")
                return []
            
            # Use original filename if provided, otherwise fall back to file path name
            file_name = original_filename if original_filename else Path(file_path).name
            
            # Generate unique document ID using the display name
            document_id = self.generate_document_id(file_name, text)
            
            # Create enhanced metadata
            file_stats = os.stat(file_path)
            metadata = {
                'source': file_path,
                'filename': file_name,  # Use the preserved original filename
                'file_type': Path(file_name).suffix.lower(),  # Get extension from original name
                'document_id': document_id,
                'content_hash': content_hash,
                'file_size': file_stats.st_size,
                'modification_time': datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
                'processing_timestamp': datetime.now().isoformat(),
                'token_count': self.get_token_count(text)
            }
            
            # Chunk the text
            documents = self.chunk_text(text, metadata)
            
            # Track processed document with original filename
            self.processed_documents[content_hash] = {
                'document_id': document_id,
                'filename': file_name,  # Store original filename
                'timestamp': datetime.now().isoformat(),
                'chunk_count': len(documents)
            }
            
            logger.info(f"Processed {file_name}: {len(documents)} chunks created with ID {document_id}")
            return documents
            
        except Exception as e:
            display_name = original_filename or Path(file_path).name
            logger.error(f"Error processing document {display_name}: {e}")
            raise
    
    def get_token_count(self, text: str) -> int:
        """Get approximate token count for text"""
        try:
            encoding = tiktoken.get_encoding("cl100k_base")
            return len(encoding.encode(text))
        except:
            # Fallback: rough estimate (4 characters per token)
            return len(text) // 4
    
    def clear_processed_cache(self):
        """Clear the processed documents cache"""
        self.processed_documents.clear()
        logger.info("Cleared processed documents cache")
    
    def get_processed_documents_info(self) -> Dict[str, Any]:
        """Get information about processed documents"""
        return self.processed_documents.copy()
    
    def remove_from_processed_cache(self, content_hash: str):
        """Remove a document from processed cache"""
        if content_hash in self.processed_documents:
            del self.processed_documents[content_hash]
            logger.info(f"Removed document with hash {content_hash} from cache")
